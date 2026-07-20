from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path(__file__).with_name("汽车焊接工业AI平台开发计划.docx")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_GRAY = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
INK = "243447"
MUTED = "667085"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_run_font(run, size=11, bold=False, color=INK, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def style_table(table, header=True):
    table.style = "Table Grid"
    for row_index, row in enumerate(table.rows):
        for cell in row.cells:
            if header and row_index == 0:
                set_cell_shading(cell, LIGHT_BLUE)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(2)
                paragraph.paragraph_format.line_spacing = 1.08
                for run in paragraph.runs:
                    set_run_font(run, size=9.3, bold=(header and row_index == 0))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header
    for row_data in rows:
        row = table.add_row()
        for index, value in enumerate(row_data):
            row.cells[index].text = str(value)
    set_table_geometry(table, widths)
    style_table(table)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    return paragraph


def add_body(doc, text, bold_lead=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_lead and text.startswith(bold_lead):
        lead = paragraph.add_run(bold_lead)
        set_run_font(lead, bold=True)
        rest = paragraph.add_run(text[len(bold_lead):])
        set_run_font(rest)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_number(doc, text, style="List Number"):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.left_indent = Inches(0.5)
    paragraph.paragraph_format.first_line_indent = Inches(-0.25)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.12
    run = paragraph.add_run(text)
    set_run_font(run)
    return paragraph


def add_callout(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [TABLE_WIDTH])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    title_run = paragraph.add_run(title + "  ")
    set_run_font(title_run, size=10.5, bold=True, color=DARK_BLUE)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, size=10.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, BLUE, 16, 8),
        "Heading 2": (13, BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for name, (size, color, before, after) in heading_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("第 ")
    set_run_font(run, size=9, color=MUTED)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    tail = paragraph.add_run(" 页")
    set_run_font(tail, size=9, color=MUTED)


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    configure_styles(doc)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run("汽车焊接工业 AI 平台 | 项目开发计划")
    set_run_font(header_run, size=8.5, color=MUTED)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(12)
    title.paragraph_format.space_after = Pt(4)
    title_run = title.add_run("汽车焊接工业 AI 平台开发计划")
    set_run_font(title_run, size=23, bold=True, color="000000")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    subtitle_run = subtitle.add_run("基于现有平台能力与 Cube Studio 功能体系的整合方案")
    set_run_font(subtitle_run, size=13.5, color=MUTED)

    metadata = [
        ("项目", "面向汽车制造过程的工业智能体关键技术开发与应用研究"),
        ("文档用途", "项目评审、开发排期与范围确认"),
        ("日期", "2026 年 7 月 12 日"),
        ("状态", "规划稿"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(2)
        label_run = paragraph.add_run(label + "：")
        set_run_font(label_run, bold=True, color="000000")
        value_run = paragraph.add_run(value)
        set_run_font(value_run, color="000000")

    doc.add_paragraph().paragraph_format.space_after = Pt(3)
    add_callout(
        doc,
        "规划结论",
        "稳定可交付版 1 个月完成，企业 MLOps 核心版累计 2-3 个月完成，全部规划功能累计 5 个月完成。",
    )

    add_heading(doc, "1. 项目概述", 1)
    add_body(
        doc,
        "当前项目已具备 React + FastAPI 的 Web 平台基础，覆盖可视化工作流、79 个算法与工业算子、基础训练、模型库、知识库、数据标注、资源监控和智能体协同。下一阶段将吸收 Cube Studio 的 MLOps、云原生算力、在线开发、模型服务与大模型能力，形成面向汽车焊接制造的工业 AI 全生命周期平台。",
    )
    add_body(
        doc,
        "建设原则：保留当前平台轻量、可视化和焊接领域专用能力，通过 Local Executor 与 Kubernetes Executor 双执行后端逐步升级，避免云原生能力阻塞现有场景交付。",
        bold_lead="建设原则：",
    )

    add_heading(doc, "2. 项目目标", 1)
    goals = [
        "在 1 个月内形成可稳定演示、可测试、可部署的工业工作流平台。",
        "在累计 2-3 个月内完成数据、实验、训练、Pipeline 调度、模型注册与推理服务闭环。",
        "在累计 5 个月内完成 Kubernetes 执行、Notebook、多集群资源、标注、AIHub、LLM 网关与工业智能体能力。",
        "建立焊接质量预测、参数推荐、异常检测和机理融合四套可复用业务模板。",
        "建立可追溯的数据、实验、模型、服务和运行链路，以及自动化测试与运维体系。",
    ]
    for goal in goals:
        add_bullet(doc, goal)

    add_heading(doc, "3. 项目范围", 1)
    add_heading(doc, "3.1 纳入范围", 2)
    in_scope = [
        "平台治理：认证、项目、角色权限、审计、通知和配置管理。",
        "数据智能：数据集、对象存储、数据探索、标注、知识库和知识图谱。",
        "工作流：可视化编排、算子、调度、重试、恢复、日志、结果与制品。",
        "模型工程：实验、训练、AutoML、模型注册、版本、评估和推理服务。",
        "云原生能力：Kubernetes 执行、Notebook、镜像、GPU 与多集群资源管理。",
        "大模型能力：RAG、智能体、统一 LLM 网关、微调模板和 AIHub。",
        "工业场景：焊接机理模型、质量预测、参数推荐、缺陷诊断与边缘部署接口。",
    ]
    for item in in_scope:
        add_bullet(doc, item)

    add_heading(doc, "3.2 暂不纳入范围", 2)
    out_scope = [
        "硬件采购、机房网络改造和第三方商业软件授权。",
        "所有国产异构算力一次性适配；首期以 NVIDIA GPU 为基线。",
        "真实商业租赁、支付和财务结算系统。",
        "与 Cube Studio 数据库、页面和内部实现保持兼容。",
    ]
    for item in out_scope:
        add_bullet(doc, item)

    add_heading(doc, "4. 当前能力与主要缺口", 1)
    capability_rows = [
        ("工作流与 DAG", "画布、保存、运行、WebSocket、条件与循环", "任务持久化、重试、恢复、调度和隔离"),
        ("算子体系", "79 个通用、机器学习、深度学习及焊接算子", "类型契约、自定义容器算子、版本与兼容性"),
        ("数据与制品", "上传、预览、JSONL DataBus、本地文件", "MinIO、数据版本、元数据、血缘和生命周期"),
        ("训练与模型", "sklearn、XGBoost、PyTorch、模型库", "实验跟踪、模型版本、审批、部署关联"),
        ("知识与标注", "RAG、向量、图谱、规则标注页面/API", "解析、重排、权限、多模态标注和审核闭环"),
        ("资源与运维", "单机指标、节点和设备信息管理", "Kubernetes、GPU 调度、配额、日志和告警"),
    ]
    add_table(doc, ["功能域", "已有能力", "主要缺口"], capability_rows, [1500, 3700, 4160])

    add_heading(doc, "5. 交付版本与周期", 1)
    release_rows = [
        ("稳定可交付版", "第 1 个月", "现有工作流、数据、算子、训练和知识库稳定化；完成工业演示闭环"),
        ("企业 MLOps 核心版", "累计第 2-3 个月", "PostgreSQL、Redis、MinIO、实验、Pipeline 调度、模型注册和推理服务"),
        ("全功能版", "累计第 5 个月", "Kubernetes、Notebook、多集群、标注、AIHub、LLM 网关和工业智能体"),
    ]
    add_table(doc, ["交付版本", "完成时间", "核心成果"], release_rows, [1850, 1500, 6010])

    add_heading(doc, "6. 分阶段开发计划", 1)
    phase_content = [
        (
            "阶段 0：基线治理与需求冻结（第 1 周）",
            [
                "建立页面、API、数据模型、算子和测试的唯一功能台账。",
                "修复构建、国际化、工作流运行、测试隔离和环境配置问题。",
                "冻结状态枚举、错误码、接口边界、数据存储和执行器架构。",
            ],
        ),
        (
            "阶段 1：稳定可交付版（第 2-4 周）",
            [
                "补齐工作流版本、运行校验、日志、重试、超时、取消和结果制品。",
                "统一 79 个算子的输入输出、异常、预览和序列化协议。",
                "完善数据集、模型库、知识库、训练和四套焊接业务模板。",
                "完成核心 E2E、部署说明、用户手册和演示脚本。",
            ],
        ),
        (
            "阶段 2：企业 MLOps 核心版（第 5-12 周）",
            [
                "迁移 PostgreSQL、Redis/Celery 和 MinIO，建立数据库迁移与制品 URI。",
                "完成实验跟踪、训练检查点、AutoML Trial、指标对比和 TensorBoard。",
                "完成定时调度、补录、重试、依赖、并发、暂停恢复和实例保留。",
                "完成模型注册、多版本、推理服务、在线测试、发布回滚、审计和通知。",
            ],
        ),
        (
            "阶段 3：云原生算力与在线开发（第 13-16 周）",
            [
                "实现 Kubernetes Job/Pod 执行器、状态同步、日志、取消和回收。",
                "实现 JupyterLab、VS Code、镜像目录、在线构建和 TensorBoard。",
                "实现集群、资源组、节点、GPU、存储挂载、配额和资源监控。",
            ],
        ),
        (
            "阶段 4：数据标注、AIHub 与大模型（第 17-20 周）",
            [
                "完成 SQL Lab、数据探索、Label Studio、多模态标注和审核闭环。",
                "生产化 RAG、检索评估、权限过滤、智能体工具与人工审核。",
                "实现 OpenAI 兼容 LLM 网关、vLLM、配额限速、调用监控和 LoRA 模板。",
                "建设模型、算子、Pipeline 和工业应用统一 AIHub。",
            ],
        ),
    ]
    for phase_title, tasks in phase_content:
        add_heading(doc, phase_title, 2)
        for task in tasks:
            add_bullet(doc, task)

    add_heading(doc, "7. 周度开发计划", 1)
    add_body(
        doc,
        "项目以周为最小管理周期。每周结束前完成代码合并、自动化测试、演示验证和周交付物归档；未通过周验收的内容不得直接转入下一阶段。",
    )

    weekly_groups = [
        (
            "第 1 月：稳定可交付版",
            [
                ("第 1 周", "基线治理与范围冻结", "盘点页面、API、模型、算子和测试；修复构建与环境问题；冻结状态、错误码、数据和执行器接口", "功能台账、技术债清单、接口基线、可构建版本"),
                ("第 2 周", "工作流与执行可靠性", "完善保存、版本、运行前校验、取消、超时、重试、日志和错误详情；修复画布交互与进度状态", "稳定工作流主链路、运行状态机、核心回归用例"),
                ("第 3 周", "数据、算子与训练闭环", "统一算子输入输出与异常协议；完善数据集、制品、训练、评估和模型保存；修复 79 个算子关键缺陷", "算子协议、数据训练闭环、稳定演示版"),
                ("第 4 周", "工业模板与首版交付", "完成质量预测、参数推荐、异常检测、机理融合模板；补齐 E2E、部署说明、用户手册和演示数据", "稳定可交付版、四套模板、首月验收报告"),
            ],
        ),
        (
            "第 2 月：企业 MLOps 核心能力建设",
            [
                ("第 5 周", "生产存储与异步任务", "迁移 PostgreSQL、Alembic、Redis/Celery 和 MinIO；统一制品 URI、配置与密钥管理", "生产数据层、对象存储、异步任务框架"),
                ("第 6 周", "实验与训练管理", "建设实验、Run、参数、指标、标签、日志和制品；接入检查点、恢复、早停、对比和 TensorBoard", "企业基础版、实验管理与训练追踪"),
                ("第 7 周", "Pipeline 调度与权限", "完成 Cron、补录、依赖、并发、超时、重试、暂停恢复；补充项目角色、审计和通知", "Pipeline 调度器、实例管理、权限审计"),
                ("第 8 周", "模型注册与基础推理", "完成模型版本、阶段、指标和审批；实现基础推理、在线测试、健康检查和服务启停", "模型注册中心、基础推理服务、第二月阶段报告"),
            ],
        ),
        (
            "第 3 月：企业 MLOps 核心版完善",
            [
                ("第 9 周", "推理服务生产化", "完善多版本发布、滚动升级、回滚、访问密钥、限流、服务日志和运行指标", "生产级推理服务、版本发布与回滚"),
                ("第 10 周", "权限、审计与通知", "完善项目角色、资源权限、关键操作审计，并接入一种企业消息通知渠道", "权限矩阵、审计日志、告警通知"),
                ("第 11 周", "系统联调与性能优化", "联调数据、Pipeline、实验、模型和服务链路；完成任务、日志、上传和推理性能优化", "全链路联调版本、性能基线和问题清单"),
                ("第 12 周", "MLOps 核心版验收", "执行功能、E2E、权限、性能、安全、备份恢复和升级验证；收敛核心缺陷", "企业 MLOps 核心版、第三月验收报告"),
            ],
        ),
        (
            "第 4 月：云原生算力与在线开发",
            [
                ("第 13 周", "Kubernetes 基础接入", "完成集群注册、命名空间、资源组、节点发现、访问凭据和连通性检查", "集群管理基础、资源发现 API"),
                ("第 14 周", "Kubernetes 执行器", "将工作流节点提交为 Job/Pod；实现状态同步、日志、取消、超时和垃圾回收", "Kubernetes Executor、任务运行闭环"),
                ("第 15 周", "Notebook、镜像与 GPU", "完成 JupyterLab、VS Code、镜像目录、在线构建、GPU 规格和节点调度", "Notebook 在线开发、镜像管理、GPU 调度"),
                ("第 16 周", "多集群与资源治理", "完成多集群路由、存储挂载、资源配额以及集群、节点、Pod 和 GPU 监控", "云原生版、多集群与资源监控"),
            ],
        ),
        (
            "第 5 月：大模型、AIHub 与总体验收",
            [
                ("第 17 周", "数据探索与标注", "建设 SQL Lab、数据探索和质量报告；对接多模态标注、审核、数据集回流和训练触发", "数据探索、标注管理与训练数据闭环"),
                ("第 18 周", "RAG 与工业智能体", "完善文档解析、Embedding、召回重排、引用和权限；完善智能体工具、重试和人工审核", "生产级 RAG、工业智能体执行闭环"),
                ("第 19 周", "LLM 网关与 AIHub", "实现模型路由、密钥、配额、限速和调用监控；统一模型、算子、Pipeline 和应用资产", "LLM 网关、AIHub 和一键开发部署"),
                ("第 20 周", "全量验收与交付", "执行全链路 E2E、性能、安全、备份恢复和升级验证；收敛缺陷；完成培训、文档和正式交付", "全功能版、验收报告、运维与用户文档"),
            ],
        ),
    ]

    for group_title, rows in weekly_groups:
        add_heading(doc, group_title, 2)
        add_table(doc, ["周次", "工作主题", "具体工作内容", "周交付物"], rows, [850, 1750, 4260, 2500])

    add_heading(doc, "8. 里程碑", 1)
    milestone_rows = [
        ("M0", "第 1 周", "基线冻结", "可构建、可测试、功能台账和接口基线完成"),
        ("M1", "第 3 周", "稳定演示版", "工作流和算子主要缺陷收敛，样例可重复运行"),
        ("M2", "第 4 周", "工业可交付版", "四套焊接模板与数据-训练-评估闭环"),
        ("M3", "第 8 周", "企业基础版", "PostgreSQL、MinIO、Redis、实验、调度和基础推理"),
        ("M4", "第 12 周", "MLOps 核心版", "生产推理、权限审计、性能优化和完整验收"),
        ("M5", "第 16 周", "云原生版", "Kubernetes、Notebook、GPU、多集群和资源管理"),
        ("M6", "第 20 周", "全功能版", "标注闭环、LLM 网关、AIHub 和工业智能体"),
    ]
    add_table(doc, ["编号", "时间", "里程碑", "验收结果"], milestone_rows, [800, 1200, 1900, 5460])

    add_heading(doc, "9. 验收标准", 1)
    acceptance = [
        "功能：用户可完成数据上传、工作流编排、训练、评估、模型注册、服务发布和结果追溯。",
        "质量：核心业务逻辑单元测试覆盖率不低于 80%，核心链路具备自动化 E2E。",
        "性能：工作流运行、文件上传、日志查询和推理接口均建立性能基线。",
        "安全：完成越权、弱密码、文件路径、自定义代码、依赖和镜像安全检查。",
        "运维：具备请求 ID、任务 ID、结构化日志、监控指标、告警、备份和升级文档。",
        "交付：提供管理员手册、用户手册、部署手册、演示数据与验收报告。",
    ]
    for item in acceptance:
        add_bullet(doc, item)

    add_heading(doc, "10. 关键风险与决策", 1)
    risk_rows = [
        ("范围持续扩大", "无法在五个月内验收", "阶段冻结范围，新增需求进入后续版本"),
        ("SQLite 与本地路径耦合", "无法扩展和迁移", "第二个月前完成 PostgreSQL 与 MinIO 迁移"),
        ("进程内 DAG 执行", "重启丢任务", "建立持久状态机、异步队列和双执行器"),
        ("云原生复杂度", "阻塞现有交付", "保留 Local Executor，Kubernetes 能力独立演进"),
        ("自定义代码安全", "任意代码执行风险", "容器隔离、权限、资源限制和审计"),
        ("异构算力范围过大", "适配成本失控", "首期以 NVIDIA GPU 为基线，其他硬件单独评估"),
    ]
    add_table(doc, ["风险", "影响", "控制措施"], risk_rows, [2200, 2500, 4660])

    add_heading(doc, "11. 待确认事项", 1)
    decisions = [
        "确认五个月范围以本文件列出的功能为上限，开发期不增加新的一级功能域。",
        "确认生产基础设施采用 PostgreSQL、Redis、MinIO 和 Kubernetes。",
        "确认首期 GPU 基线、部署环境、网络约束和外部存储方式。",
        "确认企业认证方式以及消息通知渠道。",
        "确认模型服务首期支持的模型格式和发布审批流程。",
        "确认四套焊接业务模板的样本数据、指标和验收责任人。",
    ]
    for item in decisions:
        add_number(doc, item)

    add_heading(doc, "12. 下一步行动", 1)
    next_steps = [
        "评审并冻结功能范围、五个月里程碑和验收标准。",
        "完成现有平台功能台账和技术债清单。",
        "输出目标架构、数据模型、执行器接口和部署拓扑。",
        "建立第一个月的周计划、验收用例和演示基线。",
        "从稳定工作流、存储迁移和自动化测试三条主线并行启动。",
    ]
    for item in next_steps:
        add_number(doc, item, style="List Number 2")

    add_heading(doc, "13. 周度管理机制", 1)
    weekly_management = [
        "周初计划：确认本周范围、接口变更、依赖项、验收用例和周交付物，不在周中临时增加范围。",
        "周中集成：完成跨模块接口联调，持续合并主干，及时暴露数据、权限、任务状态和部署问题。",
        "周末验收：执行自动化测试、演示验证和缺陷复盘，形成周报、版本包、测试结果和风险清单。",
        "问题升级：影响周目标的阻塞问题在发现当日升级，涉及架构和范围的决策必须记录并明确结论。",
        "版本管理：每周形成可部署版本，重要里程碑建立发布标签、数据库迁移脚本和回滚说明。",
    ]
    for item in weekly_management:
        add_bullet(doc, item)

    add_heading(doc, "14. 周交付完成标准", 1)
    done_criteria = [
        "功能代码已合并，代码审查意见已处理，无阻塞级静态检查和依赖安全问题。",
        "新增或修改功能具备单元测试、API 集成测试；主链路变化同步更新 E2E 用例。",
        "数据库、配置、接口和部署发生变化时，迁移脚本、配置样例和升级说明已同步更新。",
        "周交付物可在测试环境独立部署和演示，关键日志、指标和错误信息能够支撑问题定位。",
        "遗留问题已登记优先级、影响范围和处理计划，不以口头约定替代缺陷记录。",
    ]
    for item in done_criteria:
        add_bullet(doc, item)

    doc.core_properties.title = "汽车焊接工业 AI 平台开发计划"
    doc.core_properties.subject = "Cube Studio 功能整合与五个月开发计划"
    doc.core_properties.keywords = "汽车焊接, 工业AI, MLOps, Cube Studio, 开发计划"
    doc.core_properties.author = "项目组"
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build_document())
